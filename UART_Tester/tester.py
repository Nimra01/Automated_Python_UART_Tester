import serial    # for serial communication
import time      # for timers and delays
import struct    # for packing or unpacking a packet
import datetime
import threading   # for multithreading
from queue import Queue     # for safe transfer between threads
from docx import Document
from docx.shared import RGBColor

# ============================================================
# CRC16 (Modbus)
# ============================================================


def crc16(data: bytes):       # crc function
    crc = 0xFFFF           # initializing the crc polynomial
    for b in data:  # for loop on every bit of data
        crc ^= b     # XORing crc with every byte
        for _ in range(8):
            if crc & 1:        # if last bit is 1,
                crc = (crc >> 1) ^ 0xA001  # XOR with right shifted crc
            else:
                crc >>= 1  # if its 0, then right shift only
    return crc

# ============================================================
# Packet Generator ($FS + 11 bytes + CRC16 = 16 bytes)
# ============================================================


def create_packet(data_bytes):
    header = b'\x24\x46\x53'   # "$FS"
    body = bytes(data_bytes)  # 11 data bytes
    raw = header + body
    crc = crc16(raw)  # computing crc16
    return raw + struct.pack('<H', crc)  # little endian

# ============================================================
# SERIAL RECEIVER THREAD
# ============================================================


def serial_receiver(ser, packet_queue, stop_event):
    buffer = bytearray()  # data storing array

    while not stop_event.is_set():   # stop_event var is 0
        try:
            data = ser.read(64)  # read 64 bytes
            if not data:
                continue   # read until 64 bytes

            buffer.extend(data)

            while True:
                if len(buffer) < 16:  # check packet length
                    break

                # Check header
                if buffer[0:3] != b'\x24\x46\x53':
                    buffer.pop(0)  # drop the bytes if header is not found
                    continue

                packet = buffer[:16]

                # CRC check
                received_crc = struct.unpack('<H', packet[14:16])[0]
                calculated_crc = crc16(packet[:14])

                if received_crc != calculated_crc:
                    buffer.pop(0)  # drop the packet if crc is not same
                    continue

                # Valid packet after all 3 checks
                data_bytes = packet[3:14]  # 11 bytes
                packet_queue.put(data_bytes)  # put valid data in queue

                buffer = buffer[16:]  # all 16 bytes

        except serial.SerialException:
            break

# ============================================================
# WORD REPORT GENERATOR
# ============================================================


def generate_report(results):
    now = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"UART_Test_Report_{now}.docx"

    doc = Document()
    doc.add_heading("UART Automated Tester Report", level=1)
    doc.add_paragraph(f"Test Time: {now}")
    doc.add_paragraph(" ")

    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Parameter #"
    hdr[1].text = "Expected"
    hdr[2].text = "Received"
    hdr[3].text = "Error (%)"
    hdr[4].text = "Status"

    for r in results:
        row = table.add_row().cells
        row[0].text = str(r["id"])
        row[1].text = str(r["expected"])
        row[2].text = str(r["received"])
        row[3].text = f"{r['error']:.2f}%"

        status_cell = row[4].paragraphs[0]
        if r["status"] == "PASS":
            run = status_cell.add_run("PASS")
            run.font.color.rgb = RGBColor(0x00, 0xAA, 0x00)
        else:
            run = status_cell.add_run("FAIL")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.save(filename)
    print(f"\nReport generated: {filename}\n")

# ============================================================
# MAIN THREAD (SENDER + PROCESSOR)
# ============================================================


def main():
    ser = serial.Serial("COM6", 115200, timeout=0.1)
    time.sleep(2)

    packet_queue = Queue()  # queue used by serial thread
    stop_event = threading.Event()
    results_all = []

    # Start serial RX thread
    rx_thread = threading.Thread(
        target=serial_receiver,
        args=(ser, packet_queue, stop_event),
        daemon=True
    )
    rx_thread.start()

    print("Serial receiver thread started...")

    expected_data = [10, 20, 30, 40, 50, 60, 70, 80, 90, 100, 110]
    packets_to_send = 5

    # Send packets
    for _ in range(packets_to_send):
        packet = create_packet(expected_data)
        ser.write(packet)
        time.sleep(0.005)

    received_packets = 0

    while received_packets < packets_to_send:
        try:  # wait for the packets in queue until recv = sent packets
            recv_data = packet_queue.get(timeout=1)
        except:
            print("Timeout waiting for packet")
            continue

        received_packets += 1

        for idx in range(11):
            expected = expected_data[idx]
            received = recv_data[idx]

            error = ((received - expected) / expected) * 100
            status = "PASS" if abs(error) <= 1 else "FAIL"

            results_all.append({
                "id": idx + 1,
                "expected": expected,
                "received": received,
                "error": error,
                "status": status
            })

    stop_event.set()
    ser.close()
    generate_report(results_all)


# ============================================================
# ENTRY POINT
# ============================================================
if __name__ == "__main__":
    main()
